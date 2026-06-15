#!/usr/bin/env python3
"""Create and edit Word (.docx) files.

Usage:
    # Create from a Markdown or plain-text file
    python write.py --from-md input.md output.docx

    # Find & replace text in an existing document
    python write.py --replace "Old Text" "New Text" document.docx

    # Append a paragraph to an existing document
    python write.py --append "New paragraph text." document.docx

    # Merge multiple documents into one
    python write.py --merge doc1.docx doc2.docx doc3.docx --out merged.docx
"""

import copy
import sys


def from_markdown(src, dst):
    """Convert a Markdown / plain-text file to .docx via pandoc (preferred)
    or a simple paragraph-by-paragraph fallback."""
    import subprocess
    from pathlib import Path

    # Try pandoc first — best quality
    result = subprocess.run(["pandoc", src, "-o", dst], capture_output=True, text=True)
    if result.returncode == 0:
        print(f"Created {dst} via pandoc.")
        return

    # Fallback: plain paragraph write
    from docx import Document

    doc = Document()
    for line in Path(src).read_text(encoding="utf-8").splitlines():
        line = line.rstrip()
        if line.startswith("# "):
            doc.add_heading(line[2:], level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("### "):
            doc.add_heading(line[4:], level=3)
        elif line.startswith("- ") or line.startswith("* "):
            doc.add_paragraph(line[2:], style="List Bullet")
        elif line:
            doc.add_paragraph(line)
    doc.save(dst)
    print(f"Created {dst} (pandoc not found; used plain fallback).")


def find_replace(old, new, path):
    from docx import Document

    doc = Document(path)
    count = 0

    def replace_in_para(para):
        nonlocal count
        for run in para.runs:
            if old in run.text:
                run.text = run.text.replace(old, new)
                count += 1

    for para in doc.paragraphs:
        replace_in_para(para)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    replace_in_para(para)

    doc.save(path)
    print(f"Replaced {count} occurrence(s) of '{old}' → '{new}' in {path}")


def append_paragraph(text, path):
    from docx import Document

    doc = Document(path)
    doc.add_paragraph(text)
    doc.save(path)
    print(f"Appended paragraph to {path}")


def merge_docs(paths, out):
    from docx import Document
    from docx.oxml import OxmlElement

    merged = Document(paths[0])
    for path in paths[1:]:
        merged.add_page_break()
        sub = Document(path)
        for element in sub.element.body:
            merged.element.body.append(copy.deepcopy(element))
    merged.save(out)
    print(f"Merged {len(paths)} documents → {out}")


if __name__ == "__main__":
    args = sys.argv[1:]
    if not args or args[0] in {"-h", "--help"}:
        print(__doc__)
        sys.exit(0)

    if "--from-md" in args:
        idx = args.index("--from-md")
        src = args[idx + 1]
        dst = args[idx + 2]
        from_markdown(src, dst)

    elif "--replace" in args:
        idx = args.index("--replace")
        old = args[idx + 1]
        new = args[idx + 2]
        path = args[idx + 3]
        find_replace(old, new, path)

    elif "--append" in args:
        idx = args.index("--append")
        text = args[idx + 1]
        path = args[idx + 2]
        append_paragraph(text, path)

    elif "--merge" in args:
        idx = args.index("--merge")
        # collect paths until --out
        paths = []
        i = idx + 1
        while i < len(args) and args[i] != "--out":
            paths.append(args[i])
            i += 1
        out = args[i + 1] if "--out" in args else "merged.docx"
        merge_docs(paths, out)

    else:
        print("Unknown arguments. Run with --help for usage.")
        sys.exit(1)
