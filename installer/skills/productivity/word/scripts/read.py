#!/usr/bin/env python3
"""Read and extract content from Word (.docx) files.

Usage:
    python read.py document.docx
    python read.py document.docx --markdown
    python read.py document.docx --tables
    python read.py document.docx --metadata
    python read.py document.docx --styles
"""

import json
import sys


def extract_text(path):
    from docx import Document

    doc = Document(path)
    for para in doc.paragraphs:
        if para.text.strip():
            print(para.text)
    for i, table in enumerate(doc.tables):
        print(f"\n--- Table {i + 1} ---")
        for row in table.rows:
            print("\t".join(cell.text.strip() for cell in row.cells))


def extract_markdown(path):
    from markitdown import MarkItDown

    md = MarkItDown()
    result = md.convert(path)
    print(result.text_content)


def extract_tables(path):
    from docx import Document

    doc = Document(path)
    if not doc.tables:
        print("No tables found.")
        return
    for i, table in enumerate(doc.tables):
        print(
            f"\n--- Table {i + 1} ({len(table.rows)} rows × {len(table.columns)} cols) ---"
        )
        for row in table.rows:
            print("\t".join(cell.text.strip() for cell in row.cells))


def show_metadata(path):
    from docx import Document

    doc = Document(path)
    cp = doc.core_properties
    print(
        json.dumps(
            {
                "title": cp.title or "",
                "author": cp.author or "",
                "subject": cp.subject or "",
                "keywords": cp.keywords or "",
                "created": str(cp.created or ""),
                "modified": str(cp.modified or ""),
                "paragraphs": len(doc.paragraphs),
                "tables": len(doc.tables),
                "sections": len(doc.sections),
            },
            indent=2,
        )
    )


def show_styles(path):
    from docx import Document

    doc = Document(path)
    styles_used = sorted({p.style.name for p in doc.paragraphs if p.text.strip()})
    print("Styles in use:")
    for s in styles_used:
        print(f"  {s}")


if __name__ == "__main__":
    args = sys.argv[1:]
    if not args or args[0] in {"-h", "--help"}:
        print(__doc__)
        sys.exit(0)

    path = args[0]

    if "--metadata" in args:
        show_metadata(path)
    elif "--tables" in args:
        extract_tables(path)
    elif "--styles" in args:
        show_styles(path)
    elif "--markdown" in args:
        extract_markdown(path)
    else:
        extract_text(path)
